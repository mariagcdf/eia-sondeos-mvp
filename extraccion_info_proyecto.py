# extraccion_info_proyecto.py
from typing import Tuple, List, Dict, Any, Optional
from dotenv import load_dotenv
import json, os, re
from docx import Document

from core.extraccion.pdf_reader import (
    leer_paginas_relevantes_from_upload,
    leer_pdf_texto_completo
)
from core.extraccion.bloques_textuales import extraer_bloques_literal
from core.export_docx_template import export_docx_from_placeholder_map
from core.extraccion.regex_extract import regex_extract_min_fields

# LLM + merge
from core.extraccion.llm_utils import (
    build_prompt,
    call_llm_extract_json,
    merge_min,
)

load_dotenv(override=True)

# =============================================================
# 🔹 UTILIDADES INTERNAS
# =============================================================

def _collect_placeholders_in_doc(doc: Document) -> Dict[str, int]:
    """
    Devuelve {placeholder: cuenta} encontrados en cuerpo, tablas, encabezados y pies,
    robusto a runs troceados, NBSP y caracteres invisibles.
    """
    import regex as re2
    found: Dict[str, int] = {}

    def harvest_from_text(txt: str):
        if not txt:
            return
        txt = txt.replace("\u00A0", " ").replace("\u200B", "")
        for m in re2.finditer(r"\{\{\s*([^{}]*?)\s*\}\}", txt, flags=re2.DOTALL):
            key = (m.group(1) or "").strip()
            key = re2.sub(r"\s+", " ", key)
            if key:
                found[key] = found.get(key, 0) + 1

    def para_text(p):
        if getattr(p, "runs", None):
            return "".join(r.text for r in p.runs)
        return p.text or ""

    # cuerpo
    for p in doc.paragraphs:
        harvest_from_text(para_text(p))
    # tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    harvest_from_text(para_text(p))
    # headers / footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            harvest_from_text(para_text(p))
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        harvest_from_text(para_text(p))
        for p in section.footer.paragraphs:
            harvest_from_text(para_text(p))
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        harvest_from_text(para_text(p))
    return found


def _fmt_num(v) -> str:
    """Convierte un número en texto limpio (sin decimales innecesarios)."""
    if v is None:
        return ""
    try:
        f = float(v)
        return str(int(f)) if abs(f - int(f)) < 1e-9 else str(round(f, 3))
    except Exception:
        return str(v)


def _detectar_sondeo_existente(texto: str) -> bool:
    """Heurística simple para 'sondeo existente'."""
    return bool(re.search(r"\bsondeo\s+existente\b", texto or "", re.I))


import re as _re
def _inferir_localizacion(texto: str):
    """Fallback simple para municipio/provincia a partir del texto completo."""
    t = texto or ""
    muni = ""
    prov = ""
    m = _re.search(r"t[eé]rmino\s+municipal\s+de\s+([A-ZÁÉÍÓÚÑ][\w\s\-’'ªºáéíóúñ]+)", t, _re.I) \
        or _re.search(r"municipio\s+de\s+([A-ZÁÉÍÓÚÑ][\w\s\-’'ªºáéíóúñ]+)", t, _re.I) \
        or _re.search(r"\b([A-ZÁÉÍÓÚÑ][\w\s\-’'ªºáéíóúñ]+)\s*\(\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+\s*\)", t)
    if m:
        muni = (m.group(1) or "").strip(" ,.;:()")
    p = _re.search(r"provincia\s+de\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)", t, _re.I)
    if not p and muni:
        p = _re.search(rf"{_re.escape(muni)}\s*\(\s*([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)\s*\)", t)
    if p:
        prov = (p.group(1) or "").strip(" ,.;:()")
    return muni, prov


# =============================================================
# 🔹 1) EXTRACCIÓN DEL PDF
# =============================================================

def extract_from_project_and_eia(
    uploaded_file,
    eia_docx: Optional[str] = None,
    model: str = "gpt-4.1-mini",
    max_pages: int = 40,
    max_chars: int = 15000
) -> Tuple[Dict[str, Any], Dict[str, str], str]:
    """
    Devuelve:
      - datos_min (dict fusionado: parámetros, coordenadas, localización, particularidades)
      - literal_blocks (bloques largos como antecedentes/situación/etc.)
      - html (vista previa simple; "" si no se usa)
    """
    # 0) Lectura defensiva de "páginas relevantes" (no desempaquetamos)
    try:
        _ = leer_paginas_relevantes_from_upload(uploaded_file, max_pages=max_pages, max_chars=max_chars)
    except Exception:
        pass

    # 1) Texto completo
    try:
        texto_completo = leer_pdf_texto_completo(uploaded_file)
    except Exception:
        texto_completo = None
    if texto_completo is None:
        texto_completo = ""

    # 2) REGEX base (números/coords)
    try:
        datos_regex = regex_extract_min_fields(texto_completo)
        if not isinstance(datos_regex, dict):
            datos_regex = {}
    except Exception:
        datos_regex = {}

    # 3) Bloques literales
    try:
        literal_blocks = extraer_bloques_literal(texto_completo)
        if not isinstance(literal_blocks, dict):
            literal_blocks = {}
    except Exception:
        literal_blocks = {}

    # 🔧 Normaliza alias de bloques: usa siempre PH_Situacion (S mayúscula)
    if "PH_Situacion" not in literal_blocks and "PH_situacion" in literal_blocks:
        literal_blocks["PH_Situacion"] = literal_blocks.pop("PH_situacion")

    # 4) LLM estructurado sobre TEXTO COMPLETO + merge por campo
    try:
        prompt = build_prompt(texto_completo)
        datos_llm = call_llm_extract_json(prompt, model=model, texto_relevante=texto_completo)
    except Exception:
        datos_llm = {}

    try:
        datos_min = merge_min(datos_llm, datos_regex)
        if not isinstance(datos_min, dict):
            datos_min = {}
    except Exception:
        datos_min = datos_llm or datos_regex or {}

    # 5) Fallback muni/prov si faltan
    try:
        muni, prov = _inferir_localizacion(texto_completo)
        datos_min.setdefault("localizacion", {})
        if not (datos_min["localizacion"].get("municipio") or "").strip():
            datos_min["localizacion"]["municipio"] = muni or ""
        if not (datos_min["localizacion"].get("provincia") or "").strip():
            datos_min["localizacion"]["provincia"] = prov or ""
    except Exception:
        pass

    # 6) Consumo extra si existe
    try:
        from core.extraccion.regex_extract import extract_consumo_block
        extra = extract_consumo_block(texto_completo) or ""
        extra = re.sub(r"\s+", " ", extra).strip()
        if extra:
            literal_blocks["PH_Consumo"] = extra
    except Exception:
        pass

    # 7) Flag sondeo existente
    try:
        hay_existente = _detectar_sondeo_existente(texto_completo)
        datos_min.setdefault("flags", {})["hay_sondeo_existente"] = bool(hay_existente)
    except Exception:
        pass

    # 8) html opcional
    html = ""

    # ✅ RETURN (antes faltaba, rompía el flujo)
    return datos_min, literal_blocks, html


# =============================================================
# 🔹 2) CONSTRUCCIÓN DEL JSON DE PLACEHOLDERS
# =============================================================

import unicodedata

def _norm(s: str) -> str:
    if s is None: return ""
    s = unicodedata.normalize("NFKD", str(s))
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower().strip()
    s = re.sub(r"[\s\-]+", "_", s)
    s = re.sub(r"[^a-z0-9_\.]", "", s)
    return s

def _flatten(d: Dict[str, Any], parent_key: str = "", sep: str = ".") -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for k, v in (d or {}).items():
        k = str(k)
        key = f"{parent_key}{sep}{k}" if parent_key else k
        if isinstance(v, dict):
            out.update(_flatten(v, key, sep))
        else:
            out[key] = v
    return out

def build_placeholder_json_for_template(
    plantilla_path: str,
    datos_min: Dict[str, Any],
    literal_blocks: Dict[str, str],
    save_to: Optional[str] = None,
    incluir_plantilla: bool = True
) -> Dict[str, str]:
    """
    Construye un JSON de placeholders CANÓNICO + lo que haya en la plantilla.
    **No** añade 'localizacion.*' ni 'coordenadas.*' auxiliares.
    """
    # 1) Claves canónicas que SIEMPRE queremos en el JSON (SIN alternativas)
    canon_keys = [
        # Bloques largos
        "PH_Antecedentes", "PH_Situacion", "PH_Consumo", "geologia", "aviso_existente",
        # Coordenadas
        "utm_x_principal", "utm_y_principal", "utm_huso_principal",
        "geo_lat_principal", "geo_lon_principal",
        # Localización
        "municipio", "provincia",
        # Parámetros
        "profundidad", "diametro_inicial", "diametro_perforacion_inicial_mm",
        "caudal_max_instantaneo_l_s", "instalacion_electrica", "potencia_bombeo_kw",
    ]

    # 2) Desgloses útiles
    p   = (datos_min.get("parametros") or {})
    loc = (datos_min.get("localizacion") or {})
    c   = (datos_min.get("coordenadas") or {})
    utm = (c.get("utm") or {})
    geo = (c.get("geo") or {})
    flags = (datos_min.get("flags") or {})

    # 3) (Opcional) placeholders que existan en la plantilla
    plantilla_keys = []
    if incluir_plantilla:
        try:
            doc = Document(plantilla_path)
            from extraccion_info_proyecto import _collect_placeholders_in_doc
            plantilla_keys = list(_collect_placeholders_in_doc(doc).keys())
        except Exception:
            plantilla_keys = []

    # 4) Unión de claves objetivo
    keys_objetivo = set(canon_keys) | set(literal_blocks or {}) | set(plantilla_keys)

    # 5) Construcción del mapa canónico
    placeholder_map: Dict[str, str] = {}
    aviso_texto = (
        "⚠ Se ha detectado un SONDEO EXISTENTE en el documento. "
        "Recuerda copiar manualmente la parte referida al sondeo existente."
        if flags.get("hay_sondeo_existente") else ""
    )

    for key in keys_objetivo:
        kraw = str(key)
        val = ""

        # Bloques
        if kraw in (literal_blocks or {}):
            val = literal_blocks[kraw]

        # Aviso
        elif kraw == "aviso_existente":
            val = aviso_texto

        # Coordenadas (del merge/regex)
        elif kraw == "utm_x_principal":
            val = utm.get("x", "")
        elif kraw == "utm_y_principal":
            val = utm.get("y", "")
        elif kraw == "utm_huso_principal":
            val = utm.get("huso", "")
        elif kraw == "geo_lat_principal":
            val = geo.get("lat", "")
        elif kraw == "geo_lon_principal":
            val = geo.get("lon", "")

        # Localización (solo estos dos canónicos)
        elif kraw == "municipio":
            val = loc.get("municipio", "")
        elif kraw == "provincia":
            val = loc.get("provincia", "")

        # Parámetros
        elif kraw == "profundidad":
            val = p.get("profundidad") or p.get("profundidad_proyectada_m") or ""
        elif kraw == "diametro_inicial":
            val = p.get("diametro_inicial") or ""
        elif kraw == "diametro_perforacion_inicial_mm":
            val = p.get("diametro_perforacion_inicial_mm") or ""
        elif kraw == "caudal_max_instantaneo_l_s":
            val = p.get("caudal_max_instantaneo_l_s") or ""
        elif kraw == "instalacion_electrica":
            val = p.get("instalacion_electrica") or ""
        elif kraw == "potencia_bombeo_kw":
            val = p.get("potencia_bombeo_kw") or ""

        placeholder_map[kraw] = _fmt_num(val).replace("\u00AD", "").strip()

    # 6) Guardar si se pide
    if save_to:
        try:
            os.makedirs(os.path.dirname(save_to), exist_ok=True)
        except Exception:
            pass
        with open(save_to, "w", encoding="utf-8") as f:
            json.dump(placeholder_map, f, ensure_ascii=False, indent=2)

    return placeholder_map


# =============================================================
# 🔹 3) EXPORTACIÓN DOCX FINAL
# =============================================================

def export_eia_docx(
    datos_min: Dict[str, Any],
    literal_blocks: Dict[str, str],
    out_path: str,
    plantilla_path: str = "plantilla_EIA.docx",
    save_json_al_lado: bool = True
) -> str:
    """
    Construye el DOCX final:
      - Genera el JSON de placeholders (canónico)
      - Inyecta los datos en la plantilla
    """
    json_path = None
    if save_json_al_lado:
        out_dir = os.path.join(os.getcwd(), "resultados_pruebas")
        os.makedirs(out_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(out_path))[0]
        json_path = os.path.join(out_dir, f"{base_name}.placeholders.json")

    placeholder_map = build_placeholder_json_for_template(
        plantilla_path=plantilla_path,
        datos_min=datos_min,
        literal_blocks=literal_blocks,
        save_to=json_path
    )

    result_path = export_docx_from_placeholder_map(
        placeholder_map=placeholder_map,
        plantilla_path=plantilla_path,
        out_path=out_path
    )

    print(f"✅ Documento generado: {os.path.basename(result_path)}")
    return result_path
