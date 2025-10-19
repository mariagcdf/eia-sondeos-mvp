from typing import Dict, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from pathlib import Path
import regex as re
import subprocess
from core.sintesis.alternativas_llm import generar_alternativas_llm
import json

NBSP = "\u00A0"
ZWSP = "\u200B"
SOFT = "\u00AD"

# =====================
# Utilidades básicas
# =====================

def _clean(s: str) -> str:
    """Limpia caracteres invisibles y no imprimibles."""
    return (s or "").replace(NBSP, " ").replace(ZWSP, "").replace(SOFT, "")

def _chunks(s: str, n: int = 3000):
    """Divide un texto largo en fragmentos pequeños para evitar errores de Word."""
    for i in range(0, len(s), n):
        yield s[i:i + n]

def _para_text(p: Paragraph) -> str:
    """Extrae el texto completo de un párrafo."""
    return "".join(r.text or "" for r in getattr(p, "runs", [])) or (p.text or "")

def _clear_paragraph_keep_format(p: Paragraph):
    """Vacía un párrafo manteniendo su formato (sangría, alineación, estilo)."""
    for r in list(p.runs):
        r.text = ""
    # No eliminamos los nodos XML, solo el texto

def _insert_after(p: Paragraph, text: str = "") -> Paragraph:
    """Inserta un nuevo párrafo inmediatamente después del actual."""
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    new_para = Paragraph(new_p, p._parent)
    if text:
        for chunk in _chunks(text):
            new_para.add_run(chunk)
    return new_para

# =====================
# Escritura con saltos
# =====================

def _write_with_paragraphs(para: Paragraph, text: str):
    """
    Escribe texto respetando saltos dobles (\n\n) como nuevos párrafos,
    preservando el formato original del párrafo.
    """
    value = _clean(str(text or ""))
    blocks = [b.strip() for b in value.split("\n\n") if b.strip()]

    fmt = para.paragraph_format
    style = para.style
    alignment = para.alignment

    _clear_paragraph_keep_format(para)
    if not blocks:
        return

    para.add_run(blocks[0])
    para.paragraph_format.left_indent = fmt.left_indent
    para.paragraph_format.first_line_indent = fmt.first_line_indent
    if hasattr(fmt, "hanging_indent"):
        para.paragraph_format.hanging_indent = fmt.hanging_indent
    para.alignment = alignment
    para.style = style

    cur = para
    for b in blocks[1:]:
        cur = _insert_after(cur, b)
        cur.paragraph_format.left_indent = fmt.left_indent
        cur.paragraph_format.first_line_indent = fmt.first_line_indent
        if hasattr(fmt, "hanging_indent"):
            cur.paragraph_format.hanging_indent = fmt.hanging_indent
        cur.alignment = alignment
        cur.style = style

# =====================
# Iterador de párrafos
# =====================

def _iter_paragraphs(doc: Document):
    """Itera sobre todos los párrafos del documento (texto, tablas, encabezado y pie)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for t in section.header.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p
        for p in section.footer.paragraphs:
            yield p
        for t in section.footer.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p

# =====================
# Reemplazo de placeholders
# =====================

def _replace_placeholders(doc: Document, replacements: Dict[str, Any]):
    """
    Reemplaza todos los {{placeholders}}, incluso si están divididos por runs o contienen puntos.
    Mantiene formato y preserva los encabezados exactamente como están.
    """

    def is_in_header(para: Paragraph) -> bool:
        """Detecta si un párrafo pertenece al encabezado."""
        try:
            parent = para._element.getparent()
            while parent is not None:
                if parent.tag.endswith("hdr"):
                    return True
                parent = parent.getparent()
        except Exception:
            pass
        return False

    for para in list(_iter_paragraphs(doc)):
        full_text = _clean(_para_text(para))
        if "{{" not in full_text:
            continue

        in_header = is_in_header(para)
        replaced_any = False

        for key, val in (replacements or {}).items():
            key_pattern = re.escape(key).replace(r"\.", r"[.\s\u00A0\u200B\u00AD]*")
            pattern = re.compile(r"\{\{\s*" + key_pattern + r"\s*\}\}", re.IGNORECASE)
            if not pattern.search(full_text):
                continue

            replacement_text = str(val or "")

            # 🔹 Si el párrafo está en un encabezado, reemplazamos solo texto en runs
            if in_header:
                for run in para.runs:
                    if pattern.search(run.text):
                        run.text = pattern.sub(replacement_text, run.text)
                        replaced_any = True
                continue  # no tocar formato ni crear nuevos runs

            # 🔹 En el cuerpo normal
            if re.fullmatch(r"^\s*" + pattern.pattern + r"\s*$", full_text):
                _write_with_paragraphs(para, replacement_text)
                replaced_any = True
                break

            new_text = pattern.sub(replacement_text, full_text)
            if new_text != full_text:
                replaced_any = True
                base_run = para.runs[0] if para.runs else None
                _clear_paragraph_keep_format(para)
                new_run = para.add_run(new_text)
                if base_run:
                    font = base_run.font
                    new_run.font.name = font.name
                    new_run.font.size = font.size
                    new_run.font.bold = font.bold
                    new_run.font.italic = font.italic
                    new_run.font.underline = font.underline
                full_text = new_text

        # 🔹 No forzar mayúsculas en encabezados
        if replaced_any and not in_header:
            for r in para.runs:
                r.text = r.text


# =====================
# Relleno de tablas
# =====================

def _fill_tables_by_labels(doc: Document, data: Dict[str, Any]):
    """Rellena tablas con formato 'Etiqueta | Valor'."""
    for t in doc.tables:
        for r in t.rows:
            if len(r.cells) < 2:
                continue
            label = _clean(_para_text(r.cells[0])).lower().strip()
            if not label:
                continue
            for k, v in data.items():
                if label == k.lower().strip() and v not in (None, ""):
                    _clear_paragraph_keep_format(r.cells[1].paragraphs[0])
                    r.cells[1].paragraphs[0].add_run(str(v))
                    break

# =====================
# Función principal
# =====================

def export_docx_from_placeholder_map(
    placeholder_map: Dict[str, Any],
    plantilla_path: str,
    out_path: str,
    *,
    label_values: Optional[Dict[str, Any]] = None
) -> str:
    """
    Aplica todos los placeholders {{clave}} definidos en un JSON
    sobre una plantilla Word (.docx) y exporta el resultado.
    Mantiene formato, sangrías y estilos del párrafo original.
    """

    # 1️⃣ Ejecutar redactor automático antes de exportar
    redactor_script = Path("core/sintesis/redactar_placeholder.py")
    if redactor_script.exists():
        print("🧠 Ejecutando redactor automático de placeholders...")
        try:
            subprocess.run(["python", str(redactor_script)], check=True)
            print("✅ Redacción automática completada.")
        except subprocess.CalledProcessError as e:
            print(f"⚠️ Error ejecutando el redactor automático: {e}")
    else:
        print("⚠️ No se encontró redactar_placeholder.py — se omite la redacción automática.")

    # 2️⃣ Generar Alternativas automáticamente si faltan
    try:
        missing = [k for k in ["PH_Alternativas_Desc", "PH_Alternativas_Val", "PH_Alternativas_Just"]
                   if not placeholder_map.get(k)]
        if missing:
            print(f"🤖 Generando automáticamente las secciones de alternativas: {', '.join(missing)}")
            alt_dict = generar_alternativas_llm(placeholder_map)
            placeholder_map.update(alt_dict)

            # Guardar en el JSON más reciente para persistir los cambios
            output_dir = Path("outputs")
            json_files = sorted(output_dir.glob("*.json"), key=lambda f: f.stat().st_mtime, reverse=True)
            if json_files:
                latest_json = json_files[0]
                with open(latest_json, "r+", encoding="utf-8") as f:
                    data = json.load(f)
                    data.update(alt_dict)
                    f.seek(0)
                    json.dump(data, f, ensure_ascii=False, indent=2)
                    f.truncate()
                print(f"💾 Alternativas añadidas al JSON {latest_json.name}")
    except Exception as e:
        print(f"⚠️ Error generando alternativas automáticas: {e}")

    # 3️⃣ Procesar documento
    doc = Document(plantilla_path)

    # 🔹 Aplanar diccionario (para permitir claves con puntos)
    def flatten_dict(d, parent_key='', sep='.'):
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(flatten_dict(v, new_key, sep=sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    flat_data = flatten_dict(placeholder_map)
    replacements = {str(k): str(v or "") for k, v in flat_data.items()}

    # 🔎 Mostrar claves disponibles (debug)
    print("\n🔑 Claves disponibles para reemplazo:")
    for k in sorted(replacements.keys()):
        print(f" - {k}: {replacements[k]}")
    print("====================================\n")

    _replace_placeholders(doc, replacements)
    _fill_tables_by_labels(doc, label_values or flat_data)

    # Verificar si queda algún {{ sin reemplazar }}
    unreplaced = [p.text for p in _iter_paragraphs(doc) if "{{" in p.text]
    if unreplaced:
        print("⚠️ Placeholders sin reemplazar detectados:")
        for u in unreplaced:
            print("  -", u)
    else:
        print("✅ Todos los placeholders fueron reemplazados correctamente.")

    doc.save(out_path)
    print(f"📄 Documento exportado correctamente: {out_path}")
    return out_path
