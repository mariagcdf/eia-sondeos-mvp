# core/export_docx_template.py
from typing import Dict
import re
from docx import Document

def _chunks(s: str, n: int = 3000):
    for i in range(0, len(s), n):
        yield s[i:i+n]

def _clear_paragraph(para):
    for r in para.runs:
        r.text = ""
    while para.runs:
        para.runs[0]._element.getparent().remove(para.runs[0]._element)

def _flatten_spaces(s: str) -> str:
    s = (s or "").replace("\r", "\n").replace("\u00AD", "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()

def _write_text(para, text: str):
    for chunk in _chunks(text):
        para.add_run(chunk)

def _write_with_linebreaks(para, text: str):
    """Respeta los \n tal cual (cada \n -> w:br). No crea nuevos párrafos."""
    _clear_paragraph(para)
    text = (text or "").replace("\r", "\n")
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if part:
            _write_text(para, part)
        if i < len(parts) - 1:
            br = para.add_run(); br.add_break()

def _replace_placeholder_block_in_paragraph(para, placeholder: str, value: str) -> bool:
    token = f"{{{{{placeholder}}}}}"
    full_text = "".join(r.text for r in para.runs) if para.runs else para.text
    if token not in full_text:
        return False

    before, _, after = full_text.partition(token)

    # Si el placeholder ocupa todo el párrafo, decidimos formato:
    if before.strip() == "" and after.strip() == "":
        if placeholder in ("PH_Situacion",):  # ← SITUACIÓN: conserva saltos
            _write_with_linebreaks(para, value or "")
        # Si ya habías dejado Antecedentes con una lógica especial, puedes mantenerla aquí:
        elif placeholder == "PH_Antecedentes":
            _write_with_linebreaks(para, value or "")  # (si quieres el split antes de "El sondeo", deja tu versión anterior)
        else:
            _clear_paragraph(para)
            _write_text(para, _flatten_spaces(value or ""))
        return True

    # Caso inline: todo plano (no vamos a meter BRs dentro de líneas con más texto)
    repl = _flatten_spaces(value or "")
    new_text = full_text.replace(token, repl)
    _clear_paragraph(para)
    _write_text(para, new_text)
    return True

def _replace_placeholder_block_everywhere(doc: Document, replacements: Dict[str, str]):
    for para in doc.paragraphs:
        if "{{" in para.text:
            for k, v in replacements.items():
                _replace_placeholder_block_in_paragraph(para, k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text:
                        for k, v in replacements.items():
                            _replace_placeholder_block_in_paragraph(para, k, v)

    for section in doc.sections:
        for para in section.header.paragraphs:
            if "{{" in para.text:
                for k, v in replacements.items():
                    _replace_placeholder_block_in_paragraph(para, k, v)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{" in para.text:
                            for k, v in replacements.items():
                                _replace_placeholder_block_in_paragraph(para, k, v)
        for para in section.footer.paragraphs:
            if "{{" in para.text:
                for k, v in replacements.items():
                    _replace_placeholder_block_in_paragraph(para, k, v)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "{{" in para.text:
                            for k, v in replacements.items():
                                _replace_placeholder_block_in_paragraph(para, k, v)

def export_docx_from_placeholder_map(
    placeholder_map: Dict[str, str],
    plantilla_path: str,
    out_path: str
) -> str:
    doc = Document(plantilla_path)
    replacements = {k: ("" if v is None else str(v)) for k, v in (placeholder_map or {}).items()}
    _replace_placeholder_block_everywhere(doc, replacements)
    doc.save(out_path)
    return out_path
